from docx import Document
from docx.shared import Inches


def make_doc(customer, address, number, water_sources, name, depth=None):
    document = Document()

    document.add_heading('ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ПОДБОР ВОДООЧИСТНОГО ОБОРУДОВАНИЯ', 2)

    if depth is not None:
        water_sources_depth = f', с глубиной: {depth}'
    else:
        water_sources_depth = ''

    p = document.add_paragraph(f"""
Заказчик:                                   {customer}

Адрес объекта:                              {address}

Тел./Факс:                                  {number}

Водоисточник:                               {water_sources}{water_sources_depth}
""")
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    #
    # document.add_heading('Heading, level 1', level=2)
    # document.add_paragraph('Intense quote', style='Intense Quote')
    #
    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    # document.add_paragraph(
    #     'first item in ordered list', style='List Number'
    # )

    document.add_page_break()
    doc_name = name + '.docx'
    try:
        document.save(doc_name)
    except PermissionError:
        doc_name = '1' + doc_name
        document.save(doc_name)
    final = r"C:\Users\Григорий\PycharmProjects\Alex_bot\Alex"+f'\{doc_name}'
    return final
